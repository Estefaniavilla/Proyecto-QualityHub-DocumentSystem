import os
import time
import math
import re
from flask import Flask, render_template, request, jsonify, send_from_directory

app = Flask(__name__)

# --- Funciones de extracción de texto y metadatos ---
def extraer_texto_pdf(ruta):
    """Extrae todo el texto de un archivo PDF usando pdfplumber (sin límite de páginas)"""
    try:
        import pdfplumber
        texto = ""
        with pdfplumber.open(ruta) as pdf:
            for pagina in pdf.pages:  # Removido el límite de 10 páginas
                t = pagina.extract_text()
                if t:
                    texto += t + " "
        return texto.strip()
    except Exception:
        return ""

def extraer_texto_docx(ruta):
    """Extrae texto de un archivo Word .docx"""
    try:
        from docx import Document
        doc = Document(ruta)
        texto = " ".join([p.text for p in doc.paragraphs if p.text])
        return texto.strip()
    except Exception:
        return ""

def extraer_texto_excel(ruta):
    """Extrae texto de un archivo Excel .xls/.xlsx (sin límites)"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(ruta, read_only=True, data_only=True)
        texto = ""
        for ws in wb.worksheets:  # Removido el límite de 5 hojas
            for row in ws.iter_rows(values_only=True):  # Removido el límite de 200 filas
                for cell in row:
                    if cell is not None:
                        texto += str(cell) + " "
        wb.close()
        return texto.strip()
    except Exception:
        return ""

def extraer_texto_archivo(ruta, extension):
    """Extrae texto del archivo según su extensión"""
    ext = extension.upper()
    if ext == 'PDF':
        return extraer_texto_pdf(ruta)
    elif ext in ('DOCX', 'DOC'):
        return extraer_texto_docx(ruta)
    elif ext in ('XLS', 'XLSX'):
        return extraer_texto_excel(ruta)
    return ""

def extraer_todos_los_metadatos(ruta, extension):
    """Extrae dinámicamente absolutamente todos los metadatos disponibles del archivo y del OS"""
    metadatos = {}
    ext = extension.upper()
    try:
        # Metadatos a nivel de sistema de archivos (OS)
        stat = os.stat(ruta)
        metadatos["peso_real_bytes"] = str(stat.st_size)
        metadatos["ultima_modificacion_sistema"] = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(stat.st_mtime))
        metadatos["creacion_sistema"] = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(stat.st_ctime))
        metadatos["extension_archivo"] = extension
    except Exception as e:
        print(f"Error al extraer metadatos del OS: {e}")

    try:
        if ext == 'PDF':
            import pdfplumber
            with pdfplumber.open(ruta) as pdf:
                if pdf.metadata:
                    for k, v in pdf.metadata.items():
                        if v is not None:
                            clave_limpia = str(k).replace('.', '_').replace('$', '_')
                            metadatos[clave_limpia] = str(v)
        elif ext in ('DOCX', 'DOC'):
            from docx import Document
            doc = Document(ruta)
            props = doc.core_properties
            for prop_name in dir(props):
                if not prop_name.startswith('_'):
                    try:
                        val = getattr(props, prop_name)
                        if val is not None and not callable(val):
                            metadatos[prop_name] = str(val)
                    except Exception:
                        pass
        elif ext in ('XLS', 'XLSX'):
            from openpyxl import load_workbook
            wb = load_workbook(ruta, read_only=True)
            props = wb.properties
            for prop_name in dir(props):
                if not prop_name.startswith('_'):
                    try:
                        val = getattr(props, prop_name)
                        if val is not None and not callable(val):
                            metadatos[prop_name] = str(val)
                    except Exception:
                        pass
    except Exception as e:
        print(f"Error al extraer metadatos dinámicos: {e}")
    return metadatos

def clasificar_tipo_archivo(extension):
    """Convierte la extensión cruda a un nombre de tipo legible"""
    ext = extension.upper()
    if ext == 'PDF':
        return 'PDF'
    elif ext in ('DOC', 'DOCX'):
        return 'Word'
    elif ext in ('XLS', 'XLSX'):
        return 'Excel'
    elif ext in ('PPT', 'PPTX'):
        return 'PowerPoint'
    return ext

# Ruta interna dentro del contenedor mapeada al volumen compartido
CARPETA_ALMACENAMIENTO = '/app/storage/actuales'

try:
    from pymongo import MongoClient
    # Conectamos a la base de datos 'proyectofinal'
    cliente = MongoClient("mongodb://mongodb_container:27017/", serverSelectionTimeoutMS=5000)
    db = cliente['proyectofinal']
    coleccion_metadatos = db['metadatos']
    coleccion_usuarios = db['usuarios']
    cliente.server_info()
    print("Conexión exitosa a MongoDB (proyectofinal)")
except Exception as e:
    print(f"Error al conectar a MongoDB: {e}")
    coleccion_metadatos = None
    coleccion_usuarios = None

def obtener_version_sqlserver(archivo_descarga):
    """Intenta obtener la versión real del documento desde la base de datos de SQL Server"""
    try:
        import pymssql
        # Intentamos conectar al contenedor de SQL Server en la misma red de Docker.
        # Usualmente se llama 'sqlserver_container', 'db', 'sqlserver', o 'localhost' si corre local.
        host_options = ['sqlserver_container', 'db', 'sqlserver', 'localhost', '127.0.0.1']
        conn = None
        for host in host_options:
            try:
                # Usamos credenciales comunes de desarrollo para SQL Server en Docker
                conn = pymssql.connect(
                    server=host,
                    user='sa',
                    password='Password123!',  # Contraseña común en contenedores de desarrollo de MS SQL
                    database='proyectofinal', # Intentamos con bases de datos comunes
                    timeout=1
                )
                if conn:
                    break
            except Exception:
                continue
        
        if conn:
            cursor = conn.cursor()
            # Consultar la última versión registrada para este archivo
            query = """
                SELECT TOP 1 dv.CodigoVersion 
                FROM DocumentoVersion dv
                WHERE dv.RutaArchivo LIKE %s
                ORDER BY dv.FechaCreacion DESC
            """
            cursor.execute(query, ('%' + archivo_descarga,))
            row = cursor.fetchone()
            conn.close()
            if row and row[0]:
                version = str(row[0]).strip()
                if not version.upper().startswith('V-'):
                    version = f"V-{version}"
                return version
    except Exception as e:
        print(f"Error al conectar con SQL Server para versión: {e}")
    return None

def normalizar_fecha(fecha_str):
    """Normaliza formatos de fecha (dia/mes/año o año/mes/dia) a YYYY-MM-DD para comparación limpia"""
    if not fecha_str:
        return ""
    fecha_str = fecha_str.strip()
    if '/' in fecha_str:
        partes = fecha_str.split('/')
        if len(partes) == 3:
            if len(partes[0]) == 4:
                return f"{partes[0]}-{partes[1].zfill(2)}-{partes[2].zfill(2)}"
            elif len(partes[2]) == 4:
                return f"{partes[2]}-{partes[1].zfill(2)}-{partes[0].zfill(2)}"
    elif '-' in fecha_str:
        partes = fecha_str.split('-')
        if len(partes) == 3:
            if len(partes[0]) == 4:
                return f"{partes[0]}-{partes[1].zfill(2)}-{partes[2].zfill(2)}"
            elif len(partes[2]) == 4:
                return f"{partes[2]}-{partes[1].zfill(2)}-{partes[0].zfill(2)}"
    return fecha_str

def generar_snippet(doc_mongo, query):
    """Genera un fragmento de texto resaltando dónde coincide la búsqueda, estilo Google Search"""
    if not query:
        return ""
    
    query_lower = query.lower()
    
    # 1. Buscar coincidencia en el texto completo extraído del documento
    texto = doc_mongo.get("texto_extraido", "")
    if texto:
        idx = texto.lower().find(query_lower)
        if idx != -1:
            # Extraemos una ventana de texto alrededor de la coincidencia
            inicio = max(0, idx - 60)
            fin = min(len(texto), idx + len(query) + 60)
            snippet = texto[inicio:fin]
            
            prefijo = "... " if inicio > 0 else ""
            sufijo = " ..." if fin < len(texto) else ""
            
            # Resaltar la consulta en el fragmento usando etiquetas mark de HTML con estilo premium
            snippet_resaltado = re.sub(
                f"({re.escape(query)})", 
                r"<mark style='background-color: #fef08a; padding: 2px 4px; border-radius: 4px; color: #1e293b; font-weight: 600;'>\1</mark>", 
                snippet, 
                flags=re.IGNORECASE
            )
            return f"<strong>En el contenido del documento:</strong> {prefijo}{snippet_resaltado}{sufijo}"
    
    # 2. Buscar coincidencia en cualquiera de los otros metadatos dinámicos
    for key, val in doc_mongo.items():
        if key not in ("texto_extraido", "_id"):
            val_str = str(val)
            if query_lower in val_str.lower():
                clave_bonita = key.replace('_', ' ').upper()
                val_resaltado = re.sub(
                    f"({re.escape(query)})", 
                    r"<mark style='background-color: #fef08a; padding: 2px 4px; border-radius: 4px; color: #1e293b; font-weight: 600;'>\1</mark>", 
                    val_str, 
                    flags=re.IGNORECASE
                )
                return f"<strong>En metadato [{clave_bonita}]:</strong> {val_resaltado}"
                
    return ""

def obtener_tamano_formateado(ruta):
    """Calcula dinámicamente el peso real del archivo físico"""
    try:
        bytes_size = os.path.getsize(ruta)
        if bytes_size == 0:
            return "0 KB"
        size_name = ("B", "KB", "MB", "GB")
        i = int(math.floor(math.log(bytes_size, 1024)))
        p = math.pow(1024, i)
        s = round(bytes_size / p, 2)
        return f"{s} {size_name[i]}"
    except:
        return "Desconocido"

def mapear_formato(formato):
    """Asocia un formato del buscador a las posibles extensiones de archivo"""
    formato_upper = formato.upper()
    if formato_upper == 'WORD':
        return ['WORD', 'DOC', 'DOCX']
    elif formato_upper == 'EXCEL':
        return ['EXCEL', 'XLS', 'XLSX']
    elif formato_upper == 'POWERPOINT':
        return ['POWERPOINT', 'PPT', 'PPTX']
    elif formato_upper == 'PDF':
        return ['PDF']
    return [formato_upper]

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/login', methods=['POST'])
def login():
    data = request.json or {}
    usuario_ingresado = data.get('usuario', '').strip()
    contrasena_ingresada = data.get('contrasena', '').strip()
    compania_ingresada = data.get('compania', '').strip()
    departamento_ingresado = data.get('departamento', '').strip()

    # Si no se pudo establecer conexión con MongoDB, usamos credenciales de respaldo validadas estrictamente
    if coleccion_usuarios is None:
        if usuario_ingresado == 'fer_admin':
            if contrasena_ingresada == '12345' and compania_ingresada == 'Compania_A' and departamento_ingresado == 'Calidad':
                return jsonify({
                    "status": "success",
                    "usuario": "fer_admin",
                    "compania": "Compania_A",
                    "departamento": "Calidad"
                })
            else:
                return jsonify({"error": "Contraseña, compañía o departamento incorrectos para fer_admin"}), 401
        elif usuario_ingresado == 'auditor_global':
            if contrasena_ingresada == '54321' and compania_ingresada == 'Compania_B' and departamento_ingresado == 'Auditoría':
                return jsonify({
                    "status": "success",
                    "usuario": "auditor_global",
                    "compania": "Compania_B",
                    "departamento": "Auditoría"
                })
            else:
                return jsonify({"error": "Contraseña, compañía o departamento incorrectos para auditor_global"}), 401
        return jsonify({"error": "Usuario no registrado o datos incorrectos"}), 401

    try:
        # Buscamos en la base de datos el usuario que coincida exactamente
        usuario_doc = coleccion_usuarios.find_one({"usuario": usuario_ingresado})
        if usuario_doc:
            db_pass = str(usuario_doc.get("contrasena", "")).strip()
            db_comp = str(usuario_doc.get("compania", "")).strip()
            db_dept = str(usuario_doc.get("departamento", "")).strip()
            
            # Validación estricta cruzada
            if db_pass == contrasena_ingresada and db_comp == compania_ingresada and db_dept == departamento_ingresado:
                return jsonify({
                    "status": "success",
                    "usuario": usuario_doc["usuario"],
                    "compania": db_comp,
                    "departamento": db_dept
                })
            else:
                return jsonify({"error": "Contraseña, compañía o departamento incorrectos para este usuario"}), 401
        else:
            return jsonify({"error": "El identificador de usuario no existe"}), 401
    except Exception as e:
        return jsonify({"error": f"Error del servidor de base de datos: {str(e)}"}), 500

@app.route('/api/buscar', methods=['GET'])
def buscar_documentos():
    filtro_nombre = request.args.get('nombre', '').strip().lower()
    filtro_tipo = request.args.get('tipo', 'Todos').strip()
    filtro_fecha = request.args.get('fecha', '').strip()
    filtro_compania = request.args.get('compania', '').strip()
    filtro_departamento = request.args.get('departamento', '').strip()
    
    # Parámetros de paginación
    try:
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 5))
    except ValueError:
        page = 1
        per_page = 5

    documentos_candidatos = []

    if os.path.exists(CARPETA_ALMACENAMIENTO):
        archivos = os.listdir(CARPETA_ALMACENAMIENTO)
    else:
        archivos = []

    for archivo in archivos:
        if archivo.startswith('.'):
            continue

        ext = archivo.split('.')[-1].upper() if '.' in archivo else 'Desconocido'
        tipo_legible = clasificar_tipo_archivo(ext)
        
        # Filtro de tipo rápido previo a la base de datos
        if filtro_tipo != 'Todos':
            formatos_permitidos = mapear_formato(filtro_tipo)
            if ext not in formatos_permitidos:
                continue

        ruta_completa = os.path.join(CARPETA_ALMACENAMIENTO, archivo)

        # Consultar si ya tiene metadatos registrados en MongoDB
        doc_mongo = None
        if coleccion_metadatos is not None:
            try:
                doc_mongo = coleccion_metadatos.find_one({"archivo_descarga": archivo})
            except:
                pass

        # Obtener versión real de SQL Server para mantenerla actualizada
        version_real = obtener_version_sqlserver(archivo) or (doc_mongo.get("version") if doc_mongo else "V-1.0")

        # Extraer metadatos dinámicos
        metadatos_dinamicos = extraer_todos_los_metadatos(ruta_completa, ext)

        # Si el archivo fue subido físicamente pero no está en MongoDB, lo registramos con los metadatos del visor actual
        if not doc_mongo and coleccion_metadatos is not None:
            try:
                fecha_mod = os.path.getmtime(ruta_completa)
                fecha_real = time.strftime('%Y-%m-%d', time.localtime(fecha_mod))
            except:
                fecha_real = "2026-05-26"

            tamano_real = obtener_tamano_formateado(ruta_completa)
            nombre_limpio = archivo.rsplit('.', 1)[0].replace('_', ' ') if '.' in archivo else archivo.replace('_', ' ')
            
            # Generamos un código documental único basado en un hash estable del nombre
            codigo_hash = str(abs(hash(archivo)))[:4]
            codigo_documental = f"QH-{tipo_legible.upper()}-{codigo_hash}"

            # Extraer texto del documento para búsqueda por palabras clave
            texto_extraido = extraer_texto_archivo(ruta_completa, ext)

            # Insertamos solo los campos esenciales del sistema y unimos absolutamente todos los metadatos extraídos
            doc_mongo = {
                "nombre": nombre_limpio,
                "codigo_doc": codigo_documental,
                "version": version_real,
                "fecha": fecha_real,
                "estado": "Autorizado",
                "archivo_descarga": archivo,
                "compania_destino": filtro_compania if filtro_compania else "Compania_A",
                "departamento": filtro_departamento if filtro_departamento else "Calidad",
                "tipo": tipo_legible,
                "tamano": tamano_real,
                "texto_extraido": texto_extraido
            }
            # Guardamos todos los metadatos extraídos del archivo sin filtros
            doc_mongo.update(metadatos_dinamicos)

            try:
                coleccion_metadatos.insert_one(doc_mongo.copy())
            except:
                pass
        elif doc_mongo and coleccion_metadatos is not None:
            # Si ya existía pero queremos mantener la versión y metadatos actualizados desde SQL Server / archivo
            actualizado = False
            if doc_mongo.get("version") != version_real:
                doc_mongo["version"] = version_real
                actualizado = True
            
            # Actualizar campos dinámicos que no estén
            for k, v in metadatos_dinamicos.items():
                if doc_mongo.get(k) != v:
                    doc_mongo[k] = v
                    actualizado = True
            
            if actualizado:
                try:
                    coleccion_metadatos.replace_one({"_id": doc_mongo["_id"]}, doc_mongo)
                except:
                    pass

        # Si tenemos el documento (sea previamente guardado o recién creado en MongoDB)
        if doc_mongo:
            # 1. Filtro estricto de Compañía: El usuario solo ve los documentos asignados a su compañía
            if filtro_compania and doc_mongo.get("compania_destino", "").strip().upper() != filtro_compania.strip().upper():
                continue

            # 2. Filtro de Tipo Técnico
            if filtro_tipo != 'Todos':
                formatos_permitidos = mapear_formato(filtro_tipo)
                ext_doc = doc_mongo.get("tipo", "").upper()
                if ext_doc not in formatos_permitidos and ext not in formatos_permitidos:
                    continue

            # 3. Filtro de Fecha de Registro
            if filtro_fecha:
                fecha_doc_norm = normalizar_fecha(doc_mongo.get("fecha", ""))
                fecha_filtro_norm = normalizar_fecha(filtro_fecha)
                if fecha_doc_norm != fecha_filtro_norm:
                    continue

            # 4. Filtro por Nombre, Código, Texto Extraído o cualquier otro Metadato dinámico
            if filtro_nombre:
                nombre_doc = doc_mongo.get("nombre", "").lower()
                codigo_doc = doc_mongo.get("codigo_doc", "").lower()
                nombre_archivo = archivo.lower()
                texto_doc = doc_mongo.get("texto_extraido", "").lower()
                
                match = (filtro_nombre in nombre_doc or 
                         filtro_nombre in codigo_doc or 
                         filtro_nombre in nombre_archivo or 
                         filtro_nombre in texto_doc)

                # Si no coincide en los base, buscar en TODOS los metadatos dinámicos
                if not match:
                    for key, val in doc_mongo.items():
                        if key not in ("texto_extraido", "_id"):
                            if filtro_nombre in str(val).lower():
                                match = True
                                break

                if not match:
                    continue

            tamano_str = doc_mongo.get("tamano") or obtener_tamano_formateado(ruta_completa)

            # Generar snippet de coincidencia estilo Google (antes de descartar texto_extraido)
            snippet_match = generar_snippet(doc_mongo, filtro_nombre) if filtro_nombre else ""

            # Construir el candidato conservando todos los metadatos dinámicos
            doc_candidato = dict(doc_mongo)
            doc_candidato.pop("_id", None)
            doc_candidato.pop("texto_extraido", None)

            # Inyectar el snippet de coincidencia en el resultado
            doc_candidato["snippet_match"] = snippet_match
            
            # Asegurar compatibilidad de campos básicos para el renderizado del listado
            if "archivo_descarga" not in doc_candidato:
                doc_candidato["archivo_descarga"] = archivo
            if "nombre" not in doc_candidato:
                doc_candidato["nombre"] = doc_mongo.get("nombre", archivo)
            if "fecha" not in doc_candidato:
                doc_candidato["fecha"] = doc_mongo.get("fecha", "2026-05-26")
            if "estado" not in doc_candidato:
                doc_candidato["estado"] = doc_mongo.get("estado", "Autorizado")
            if "tipo" not in doc_candidato:
                doc_candidato["tipo"] = tipo_legible
            if "version" not in doc_candidato:
                doc_candidato["version"] = version_real

            documentos_candidatos.append(doc_candidato)

    # Calcular los conteos de KPI sobre el conjunto TOTAL filtrado antes de paginar
    conteo_autorizados = sum(1 for doc in documentos_candidatos if doc["estado"] in ['Authorized', 'Autorizado'])
    conteo_revision = len(documentos_candidatos) - conteo_autorizados

    # Calcular Paginación
    total_items = len(documentos_candidatos)
    total_pages = math.ceil(total_items / per_page) if total_items > 0 else 1
    
    if page < 1:
        page = 1
    elif page > total_pages:
        page = total_pages

    start_idx = (page - 1) * per_page
    end_idx = start_idx + per_page
    paginated_docs = documentos_candidatos[start_idx:end_idx]

    return jsonify({
        "documentos": paginated_docs,
        "total": total_items,
        "page": page,
        "pages": total_pages,
        "per_page": per_page,
        "kpis": {
            "totales": total_items,
            "autorizados": conteo_autorizados,
            "revision": conteo_revision
        }
    })

@app.route('/api/metadatos/<path:nombre_archivo>', methods=['GET'])
def obtener_metadatos(nombre_archivo):
    doc = None
    if coleccion_metadatos is not None:
        try:
            doc = coleccion_metadatos.find_one({"archivo_descarga": nombre_archivo})
        except:
            pass
    
    if doc:
        # Retornamos el documento completo de MongoDB (que contiene todos los metadatos)
        # Convertimos _id a string y removemos texto_extraido para no sobrecargar
        res_data = dict(doc)
        res_data.pop("_id", None)
        res_data.pop("texto_extraido", None)
        return jsonify(res_data)
    
    return jsonify({"error": "No se encontraron metadatos"}), 404

@app.route('/api/metadatos/editar', methods=['POST'])
def editar_metadatos():
    data = request.json or {}
    archivo_descarga = data.get("archivo_descarga")
    
    if not archivo_descarga:
        return jsonify({"error": "Identificador de archivo faltante"}), 400
        
    estado_actual = "En Revisión (No)"
    doc_existente = {}
    if coleccion_metadatos is not None:
        try:
            temp = coleccion_metadatos.find_one({"archivo_descarga": archivo_descarga})
            if temp:
                doc_existente = dict(temp)
                estado_actual = doc_existente.get("estado", "En Revisión (No)")
        except:
            pass

    # Mantener metadatos dinámicos existentes
    valores_actualizados = doc_existente.copy()
    valores_actualizados.pop("_id", None)
    
    # Actualizar con los nuevos valores del POST
    valores_actualizados.update({
        "nombre": data.get("nombre"),
        "compania_destino": data.get("compania_destino"),
        "departamento": data.get("departamento"),
        "estado": estado_actual,
        "archivo_descarga": archivo_descarga,
        "fecha": data.get("fecha"),
        "codigo_doc": data.get("codigo_doc", "QH-GEN-001"),
        "version": data.get("version", "V-1.0"),
        "autor_subida": data.get("autor_subida", "Estefanía (Gestor Documental)"),
        "tamano": data.get("tamano"),
        "clasificacion": data.get("clasificacion", "Uso Interno")
    })
    
    if coleccion_metadatos is not None:
        coleccion_metadatos.update_one(
            {"archivo_descarga": archivo_descarga},
            {"$set": valores_actualizados},
            upsert=True
        )
    
    return jsonify({"status": "success", "message": "Metadatos actualizados con éxito"})

@app.route('/api/ver/<path:nombre_archivo>')
def ver_archivo_inline(nombre_archivo):
    """Envía el archivo con cabeceras de visualización inline para forzar su apertura en el navegador"""
    response = send_from_directory(CARPETA_ALMACENAMIENTO, nombre_archivo, as_attachment=False)
    if nombre_archivo.lower().endswith('.pdf'):
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = 'inline'
    elif nombre_archivo.lower().endswith('.docx'):
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif nombre_archivo.lower().endswith('.xls') or nombre_archivo.lower().endswith('.xlsx'):
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    return response

@app.route('/api/descargar/<path:nombre_archivo>')
def descargar_archivo_attachment(nombre_archivo):
    """Fuerza la descarga directa del archivo físico de la carpeta storage"""
    return send_from_directory(CARPETA_ALMACENAMIENTO, nombre_archivo, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)